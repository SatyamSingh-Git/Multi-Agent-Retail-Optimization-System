# main.py
import os
import json
from typing import TypedDict, List, Tuple, Dict, Any
import sqlite3
import pandas as pd
import datetime # Needed for timestamp in filename
from docx import Document

from langgraph.graph import StateGraph, END

# Import agent functions
from agents.data_ingestion_agent import run_data_ingestion
from agents.demand_forecaster_agent import run_demand_forecasting
from agents.inventory_monitoring_agent import run_inventory_monitoring
# --- NEW IMPORTS ---
from agents.replenishment_agent import run_replenishment_calculation
from agents.pricing_optimizer_agent import run_pricing_optimization
from agents.supplier_agent import run_supplier_interaction
from utils.ollama_utils import get_ollama_completion # Need LLM here

# Define the state structure for the graph
class AgentState(TypedDict):
    # Input/Control
    items_to_process: List[Tuple[int, int]] | None

    # Agent Outputs
    data_ingestion_status: str | None
    forecast_results: Dict | None
    inventory_status: Dict | None
    # --- NEW STATE FIELDS ---
    replenishment_proposals: Dict | None # Output from replenishment agent
    pricing_proposals: Dict | None       # Output from pricing agent
    order_placement_results: Dict | None # Output from supplier agent

    # General Status
    error_message: str | None


# Database path needed for query
DB_PATH = os.path.join(os.path.dirname(__file__), 'database', 'retail_data.db')

def get_existing_items(limit=15, history_days=90):
    """
    Queries DB for items in BOTH inventory & recent demand history.
    Uses explicit DATE() casting for reliable comparison in SQLite.
    """
    items = []
    conn = None
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        # --- **** MODIFIED QUERY with explicit DATE() casting **** ---
        query = f"""
            SELECT DISTINCT i.ProductID, i.StoreID
            FROM inventory_monitoring i
            INNER JOIN demand_forecast d ON i.ProductID = d.ProductID AND i.StoreID = d.StoreID
            WHERE DATE(d.Date) >= DATE('now', '-{history_days} days')
            LIMIT ?
        """
        # --- **** END MODIFIED QUERY **** ---

        print(f"DEBUG GET_ITEMS: Executing SQL:\n{query}\nWith limit={limit}, history_days={history_days}") # Keep this debug print
        cursor.execute(query, (limit,))
        items = cursor.fetchall()
        print(f"DEBUG GET_ITEMS: Raw fetchall result from DB query: {items}") # Keep this debug print
        print(f"Retrieved {len(items)} existing Product/Store pairs present in both inventory and recent demand history.")
    except sqlite3.Error as e:
        print(f"Error querying existing items from DB: {e}")
        # Avoid using st.error directly in backend logic if possible
        # Log the error or raise it to be caught by the Streamlit part
    except Exception as e:
        print(f"Unexpected error in get_existing_items: {e}")
    finally:
        if conn: conn.close()
    return items


# --- Define Agent Nodes (with State Update Fix) ---

def data_ingestion_node(state: AgentState) -> AgentState:
    """Node for data ingestion. Returns full updated state."""
    print("\n--- (1) Executing Data Ingestion Node ---")
    current_state = state.copy()
    try:
        result = run_data_ingestion()
        status = result.get("status", "Unknown Status")
        print(f"Data Ingestion Result: {status}")
        current_state["data_ingestion_status"] = status
        if "Failed" in status or "Partial" in status:
             error_msg = f"Data ingestion failed or partially failed. Details: {result.get('details')}"
             print(f"ERROR: {error_msg}")
             current_state["error_message"] = error_msg
             current_state["items_to_process"] = []
             return current_state
        else:
             items_from_db = get_existing_items(limit=5)
             if not items_from_db:
                 error_msg = "Ingestion successful, but NO valid Product/Store pairs found in BOTH inventory and recent demand history."
                 print(f"WARNING: {error_msg}")
                 current_state["items_to_process"] = []
                 current_state["error_message"] = error_msg
                 return current_state
             print(f"Ingestion successful. Using items found in DB for next steps: {items_from_db}")
             current_state["items_to_process"] = items_from_db
             current_state["error_message"] = None
             return current_state
    except Exception as e:
        print(f"Error during data ingestion node execution: {e}")
        import traceback
        error_details = traceback.format_exc()
        current_state["data_ingestion_status"] = "Execution Error"
        current_state["error_message"] = f"Node failed: {e}\n{error_details}"
        return current_state


def demand_forecast_node(state: AgentState) -> AgentState:
    """Node for demand forecasting. Returns full updated state."""
    print("\n--- (2) Executing Demand Forecast Node ---")
    current_state = state.copy()
    items = current_state.get("items_to_process")
    if not items:
        print("No items specified for forecasting. Skipping.")
        current_state["forecast_results"] = {"status": "Skipped", "message": "No items to process"}
        return current_state
    print(f"Running forecasts for {len(items)} item(s)...")
    try:
        results = run_demand_forecasting(items)
        print(f"Forecasting Node Summary: Processed: {results.get('processed_count')}, Succeeded: {results.get('success_count')}")
        current_state["forecast_results"] = results
        return current_state
    except Exception as e:
        print(f"Error during demand forecast node execution: {e}")
        import traceback
        error_details = traceback.format_exc()
        current_state["forecast_results"] = {"status": "Execution Error"}
        current_state["error_message"] = f"Node failed: {e}\n{error_details}"
        return current_state


def inventory_monitor_node(state: AgentState) -> AgentState:
    """Node for inventory monitoring. Returns full updated state."""
    print("\n--- (3) Executing Inventory Monitor Node ---")
    current_state = state.copy()
    items = current_state.get("items_to_process")
    if not items:
        print("No items specified for monitoring. Skipping.")
        current_state["inventory_status"] = {"status": "Skipped", "message": "No items to process"}
        return current_state
    print(f"Running inventory monitoring for {len(items)} item(s)...")
    try:
        results = run_inventory_monitoring(items)
        summary = results.get('summary', {})
        print(f"Monitoring Node Summary: Processed: {summary.get('processed_count', 'N/A')}")
        current_state["inventory_status"] = results
        return current_state
    except Exception as e:
        print(f"Error during inventory monitor node execution: {e}")
        import traceback
        error_details = traceback.format_exc()
        current_state["inventory_status"] = {"status": "Execution Error"}
        current_state["error_message"] = f"Node failed: {e}\n{error_details}"
        return current_state

def replenishment_node(state: AgentState) -> AgentState:
    """Node for replenishment calculation. Returns full updated state."""
    print("\n--- (4) Executing Replenishment Node ---") # Changed numbering
    current_state = state.copy()
    inv_status = current_state.get("inventory_status")

    if not inv_status or inv_status.get("status") == "Skipped" or not isinstance(inv_status, dict):
        print("Skipping replenishment: No valid inventory status available.")
        current_state["replenishment_proposals"] = {"status": "Skipped", "message": "Inventory status missing or invalid.", "proposed_orders": []}
        return current_state

    print(f"Running replenishment calculations based on inventory status...")
    try:
        status_to_process = {k: v for k, v in inv_status.items() if k != 'summary'}
        if not status_to_process:
             print("No inventory items found in status dictionary to process for replenishment.")
             current_state["replenishment_proposals"] = {"status": "Skipped", "message": "No items in inventory status.", "proposed_orders": []}
             return current_state

        results = run_replenishment_calculation(status_to_process)
        print(f"Replenishment Node Summary: Processed: {results.get('processed_count')}, Proposed Orders: {results.get('proposal_count')}")
        current_state["replenishment_proposals"] = results
        return current_state
    except Exception as e:
        print(f"Error during replenishment node execution: {e}")
        import traceback
        error_details = traceback.format_exc()
        current_state["replenishment_proposals"] = {"status": "Execution Error", "proposed_orders": []}
        current_state["error_message"] = f"Node failed: {e}\n{error_details}"
        return current_state

def pricing_optimization_node(state: AgentState) -> AgentState:
    """Node for pricing optimization. Returns full updated state."""
    print("\n--- (6) Executing Pricing Optimization Node ---") # Changed numbering
    current_state = state.copy()
    inv_status = current_state.get("inventory_status")

    if not inv_status or inv_status.get("status") == "Skipped" or not isinstance(inv_status, dict):
        print("Skipping pricing optimization: No valid inventory status available.")
        current_state["pricing_proposals"] = {"status": "Skipped", "message": "Inventory status missing or invalid.", "proposed_actions": []}
        return current_state

    print(f"Running pricing optimization based on inventory status...")
    try:
        status_to_process = {k: v for k, v in inv_status.items() if k != 'summary'}
        if not status_to_process:
             print("No inventory items found in status dictionary to process for pricing.")
             current_state["pricing_proposals"] = {"status": "Skipped", "message": "No items in inventory status.", "proposed_actions": []}
             return current_state

        results = run_pricing_optimization(status_to_process)
        print(f"Pricing Node Summary: Processed: {results.get('processed_count')}, Proposed Actions: {results.get('proposal_count')}")
        current_state["pricing_proposals"] = results
        return current_state
    except Exception as e:
        print(f"Error during pricing optimization node execution: {e}")
        import traceback
        error_details = traceback.format_exc()
        current_state["pricing_proposals"] = {"status": "Execution Error", "proposed_actions": []}
        current_state["error_message"] = f"Node failed: {e}\n{error_details}"
        return current_state

def supplier_interaction_node(state: AgentState) -> AgentState:
    """Node for supplier interaction (simulated order placement)."""
    print("\n--- (5) Executing Supplier Interaction Node ---") # Changed numbering
    current_state = state.copy()
    repl_proposals = current_state.get("replenishment_proposals")
    proposed_orders = []

    # Ensure repl_proposals is a dict and get proposed_orders safely
    if isinstance(repl_proposals, dict):
        proposed_orders = repl_proposals.get("proposed_orders", [])

    if not proposed_orders:
         print("No proposed orders found to send to supplier. Skipping.")
         current_state["order_placement_results"] = {"status": "Skipped", "message": "No orders proposed.", "placed_orders": []}
         return current_state

    print(f"Simulating order placement for {len(proposed_orders)} proposal(s)...")
    try:
        results = run_supplier_interaction(proposed_orders)
        print(f"Supplier Node Summary: Processed: {results.get('processed_count')}, Orders Placed: {results.get('placed_count')}")
        current_state["order_placement_results"] = results
        return current_state
    except Exception as e:
        print(f"Error during supplier interaction node execution: {e}")
        import traceback
        error_details = traceback.format_exc()
        current_state["order_placement_results"] = {"status": "Execution Error", "placed_orders": []}
        current_state["error_message"] = f"Node failed: {e}\n{error_details}"
        return current_state


# --- *** NEW: Conflict Resolution Node *** ---
def resolve_conflicts_node(state: AgentState) -> AgentState:
    """
    Analyzes proposals and resolves potential conflicts before final actions.
    Example: Check if total proposed orders exceed capacity (if capacity is available).
    Uses LLM for complex trade-offs if needed.
    """
    print("\n--- (X) Executing Conflict Resolution Node ---")
    current_state = state.copy()
    repl_proposals_dict = current_state.get("replenishment_proposals", {})
    pricing_proposals_dict = current_state.get("pricing_proposals", {})
    inventory_status_dict = current_state.get("inventory_status", {})

    proposed_orders = repl_proposals_dict.get("proposed_orders", []) if isinstance(repl_proposals_dict, dict) else []
    proposed_pricing = pricing_proposals_dict.get("proposed_actions", []) if isinstance(pricing_proposals_dict, dict) else []

    final_orders = list(proposed_orders) # Start with existing proposals
    final_pricing = list(proposed_pricing)
    resolution_log = []

    # --- Example Conflict 1: Simple Capacity Check (Needs Warehouse Info) ---
    # This is a placeholder as we don't have total warehouse capacity easily available yet.
    # total_order_qty = sum(p.get('QuantityOrdered', 0) for p in final_orders)
    # total_warehouse_capacity = get_total_warehouse_capacity() # Needs implementation
    # if total_order_qty > total_warehouse_capacity:
    #     resolution_log.append(f"WARN: Total proposed order qty ({total_order_qty}) exceeds capacity ({total_warehouse_capacity}). Needs prioritization.")
        # --- Add prioritization logic here (rule-based or LLM) ---
        # Example Rule: Prioritize items most below ROP or with highest forecast
        # Example LLM:
        # prompt = f"Orders {final_orders} exceed capacity {total_warehouse_capacity}. Prioritize based on stock levels relative to reorder point and potential stockouts. Adjust quantities."
        # adjusted_orders_str = get_ollama_completion(prompt)
        # final_orders = parse_adjusted_orders(adjusted_orders_str) # Needs implementation

    # --- Example Conflict 2: Pricing vs. Replenishment (Simple Rule) ---
    items_needing_order = {p['ProductID'] for p in final_orders}
    items_getting_discount = {p['ProductID'] for p in final_pricing if p['RecommendedPrice'] < p['CurrentPrice']}

    conflicting_items = items_needing_order.intersection(items_getting_discount)
    if conflicting_items:
        msg = f"Potential Conflict: Items {conflicting_items} need replenishment AND are proposed for discount."
        resolution_log.append(msg)
        print(f"  {msg}")
        # Simple Resolution: Maybe delay discount if stock is critically low?
        # For now, just log it. More complex logic could modify final_pricing.

    print("  Conflict Resolution Log:", resolution_log if resolution_log else "No conflicts detected.")

    # Update state with potentially modified proposals
    # Ensure the structure matches what the supplier agent expects
    current_state["replenishment_proposals"]["proposed_orders"] = final_orders # Update orders list

    # We might create a new key for 'final_pricing_actions' if needed
    # current_state["final_pricing_actions"] = final_pricing

    return current_state


# --- Build the Graph (Modify Edges) ---
def build_graph():
    """Builds the LangGraph state machine with conflict resolution."""
    workflow = StateGraph(AgentState)

    # Add nodes (Existing + New)
    workflow.add_node("data_ingestion", data_ingestion_node)
    workflow.add_node("demand_forecast", demand_forecast_node)
    workflow.add_node("inventory_monitor", inventory_monitor_node)
    workflow.add_node("replenishment_calc", replenishment_node)
    workflow.add_node("pricing_optimiser", pricing_optimization_node)
    # --- Add Conflict Node ---
    workflow.add_node("resolve_conflicts", resolve_conflicts_node)
    workflow.add_node("supplier_interact", supplier_interaction_node)

    # Define edges
    workflow.set_entry_point("data_ingestion")
    workflow.add_edge("data_ingestion", "demand_forecast")
    workflow.add_edge("demand_forecast", "inventory_monitor")

    # --- Modified Flow ---
    # After monitoring, run replenishment and pricing (sequentially for simplicity now, avoid Send issues)
    workflow.add_edge("inventory_monitor", "replenishment_calc")
    workflow.add_edge("replenishment_calc", "pricing_optimiser")

    # After both action proposals are generated, resolve conflicts
    workflow.add_edge("pricing_optimiser", "resolve_conflicts")

    # After resolving conflicts, interact with supplier (using potentially adjusted orders)
    workflow.add_edge("resolve_conflicts", "supplier_interact")

    # End after supplier interaction
    workflow.add_edge("supplier_interact", END)


    app = workflow.compile()
    print("Graph compiled successfully with Phase 4 nodes (incl. Conflict Resolution).")
    return app

# --- NEW FUNCTION: Generate Word Document Report ---
def generate_docx_report(final_state: AgentState, filename: str):
    """Generates a formatted .docx report from the final agent state."""
    print(f"\n--- Generating Word Report: {filename} ---")
    try:
        document = Document()
        document.add_heading('AI Agent Workflow Summary Report', level=0)

        run_timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        document.add_paragraph(f"Report Generated: {run_timestamp}")
        document.add_paragraph() # Add a blank line

        if not final_state:
            document.add_paragraph("Workflow execution did not produce a final state.")
            document.save(filename)
            print(f"Report saved (empty state).")
            return

        # --- Overall Status ---
        document.add_heading('Overall Status', level=1)
        error_msg = final_state.get("error_message")
        ingestion_status = final_state.get("data_ingestion_status", "Unknown")
        items_list = final_state.get('items_to_process', []) # Default to empty list
        items_processed_flag = isinstance(items_list, list) and len(items_list) > 0

        status_para = document.add_paragraph()
        status_para.add_run("Overall Status: ").bold = True
        if error_msg and ingestion_status != "Success":
            status_para.add_run("CRITICAL ERROR").bold = True
        elif error_msg:
             status_para.add_run("Completed with Warnings/Errors").bold = True
        elif not items_processed_flag and ingestion_status == "Success":
             status_para.add_run("Completed (No items required processing)").bold = True
        else:
            status_para.add_run("Completed Successfully").bold = True

        # --- Data & Setup ---
        document.add_heading('Data & Setup', level=1)
        p_ingest = document.add_paragraph()
        p_ingest.add_run("Data Ingestion Status: ").bold = True
        p_ingest.add_run(str(ingestion_status))

        items_count = len(items_list) if isinstance(items_list, list) else 0
        p_items = document.add_paragraph()
        p_items.add_run("Items Selected for Processing: ").bold = True
        p_items.add_run(str(items_count))
        if items_count > 0 and isinstance(items_list, list):
             # Show the actual items processed
             items_str = ", ".join([f"P:{p}/S:{s}" for p, s in items_list])
             document.add_paragraph(f"Processed Pairs: {items_str}", style='List Bullet')


        # --- Analysis Results ---
        if items_count > 0: # Only show analysis if items were processed
            document.add_heading('Analysis Results', level=1)
            fc_res = final_state.get('forecast_results', {})
            fc_proc = fc_res.get('processed_count', 0)
            fc_succ = fc_res.get('success_count', 0)
            p_fc = document.add_paragraph()
            p_fc.add_run("Demand Forecasting: ").bold = True
            p_fc.add_run(f"Processed={fc_proc}, Succeeded={fc_succ}")

            inv_res = final_state.get('inventory_status', {})
            inv_proc = inv_res.get('summary', {}).get('processed_count', 0)
            # Count specific flags
            low_stock_count = 0
            excess_stock_count = 0
            near_expiry_count = 0
            expired_count = 0
            if isinstance(inv_res, dict):
                for key, value in inv_res.items():
                    if key != 'summary' and isinstance(value, dict):
                        flags = value.get('flags', [])
                        if 'LOW_STOCK' in flags: low_stock_count += 1
                        if 'EXCESS_STOCK' in flags: excess_stock_count += 1
                        if 'NEAR_EXPIRY' in flags: near_expiry_count += 1
                        if 'EXPIRED' in flags: expired_count += 1

            p_inv = document.add_paragraph()
            p_inv.add_run("Inventory Monitoring: ").bold = True
            p_inv.add_run(f"Processed={inv_proc}")
            # Add flags as bullet points
            document.add_paragraph(f"Items Flagged Low Stock:   {low_stock_count}", style='List Bullet')
            document.add_paragraph(f"Items Flagged Excess Stock: {excess_stock_count}", style='List Bullet')
            document.add_paragraph(f"Items Flagged Near Expiry: {near_expiry_count}", style='List Bullet')
            document.add_paragraph(f"Items Flagged Expired:     {expired_count}", style='List Bullet')


            # --- Action Summary ---
            document.add_heading('Actions Proposed / Taken', level=1)
            rep_res = final_state.get('replenishment_proposals', {})
            rep_prop_count = rep_res.get('proposal_count', 0)
            p_rep = document.add_paragraph()
            p_rep.add_run("Replenishment Orders Proposed: ").bold = True
            p_rep.add_run(str(rep_prop_count))
            if rep_prop_count > 0:
                document.add_paragraph("Examples:")
                for order in rep_res.get('proposed_orders', [])[:5]: # Show up to 5 examples
                    document.add_paragraph(f"P:{order.get('ProductID')} S:{order.get('StoreID')} Qty:{order.get('QuantityOrdered')}", style='List Bullet')

            sup_res = final_state.get('order_placement_results', {})
            sup_placed_count = sup_res.get('placed_count', 0)
            p_sup = document.add_paragraph()
            p_sup.add_run("Orders Placed (Simulated): ").bold = True
            p_sup.add_run(str(sup_placed_count))
            if sup_placed_count > 0:
                 document.add_paragraph("Examples:")
                 for order in sup_res.get('placed_orders', [])[:5]: # Show up to 5 examples
                     document.add_paragraph(f"ID:{order.get('OrderID')} P:{order.get('ProductID')} S:{order.get('StoreID')} Qty:{order.get('QuantityOrdered')} ETA:{order.get('ExpectedDeliveryDate')}", style='List Bullet')

            prc_res = final_state.get('pricing_proposals', {})
            prc_prop_count = prc_res.get('proposal_count', 0)
            p_prc = document.add_paragraph()
            p_prc.add_run("Pricing Actions Proposed: ").bold = True
            p_prc.add_run(str(prc_prop_count))
            if prc_prop_count > 0:
                 document.add_paragraph("Examples:")
                 for action in prc_res.get('proposed_actions', [])[:5]: # Show up to 5 examples
                      document.add_paragraph(f"P:{action.get('ProductID')} S:{action.get('StoreID')} New Price:{action.get('RecommendedPrice')} (Reason: {action.get('Reason')})", style='List Bullet')

        # --- Errors / Warnings ---
        if error_msg:
            document.add_heading('Warnings / Errors Logged', level=1)
            # Print only the message part for clarity
            error_lines = str(error_msg).split('\n')
            document.add_paragraph(f"- {error_lines[0]}") # Print first line or the whole thing if one line


        # --- Save Document ---
        document.save(filename)
        print(f"Report successfully saved to {filename}")

    except ImportError:
        print("Error: 'python-docx' library not found. Cannot generate Word report.")
        print("Please install it using: pip install python-docx")
    except Exception as e:
        print(f"Error generating Word report: {e}")
        import traceback
        traceback.print_exc()

# --- Main Execution ---
if __name__ == "__main__":
    print("Starting Retail Inventory Agent Workflow (Phase 4)...")

    # Database checks...
    db_file = os.path.join(os.path.dirname(__file__), 'database', 'retail_data.db')
    if not os.path.exists(db_file):
         print("ERROR: Database file not found. Please run database/database_setup.py")
         exit()
    else:
         # Check for orders table (optional but good)
         conn_check = sqlite3.connect(db_file)
         cursor_check = conn_check.cursor()
         try:
             cursor_check.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='orders';")
             if not cursor_check.fetchone():
                  print("WARNING: 'orders' table missing. Run database/database_setup.py.")
         except Exception as e: print(f"DB check error: {e}")
         finally:
              cursor_check.close()
              conn_check.close()

    app = build_graph() # Uses the sequential graph build function

    print("\n--- Running Graph (Phase 4) ---")
    initial_state: AgentState = {
        "items_to_process": None, "data_ingestion_status": None,
        "forecast_results": None, "inventory_status": None,
        "replenishment_proposals": None, "pricing_proposals": None,
        "order_placement_results": None, "error_message": None
        }

    print("\nStreaming Graph Execution Steps:")
    final_state = None
    for step_output in app.stream(initial_state):
        step_name = list(step_output.keys())[0]
        step_state = step_output[step_name]
        print(f"\n--- State after Step: {step_name} ---")

        # Safer State Printing During Stream (Keep this for debugging if needed)
        items_proc = step_state.get('items_to_process')
        print(f"  Items to Process: {items_proc if items_proc is not None else 'Not Set'}")
        print(f"  Ingestion Status: {step_state.get('data_ingestion_status')}")
        forecast_res = step_state.get('forecast_results')
        forecast_summary_str = "N/A"
        if isinstance(forecast_res, dict):
            forecast_summary_str = f"{forecast_res.get('success_count', 'N/A')} successful forecasts"
        print(f"  Forecast Result Summary: {forecast_summary_str}")
        inventory_stat = step_state.get('inventory_status')
        inventory_summary_str = "N/A"
        if isinstance(inventory_stat, dict):
            inventory_summary = inventory_stat.get('summary', {})
            inventory_summary_str = f"{inventory_summary.get('processed_count', 'N/A')} items monitored"
        print(f"  Inventory Status Summary: {inventory_summary_str}")
        repl_props = step_state.get('replenishment_proposals')
        repl_summary_str = "N/A"
        if isinstance(repl_props, dict):
            repl_summary_str = f"{repl_props.get('proposal_count', 'N/A')} proposed orders"
        print(f"  Replenishment Summary: {repl_summary_str}")
        price_props = step_state.get('pricing_proposals')
        price_summary_str = "N/A"
        if isinstance(price_props, dict):
            price_summary_str = f"{price_props.get('proposal_count', 'N/A')} proposed actions"
        print(f"  Pricing Summary: {price_summary_str}")
        order_res = step_state.get('order_placement_results')
        order_summary_str = "N/A"
        if isinstance(order_res, dict):
            order_summary_str = f"{order_res.get('placed_count', 'N/A')} orders placed"
        print(f"  Order Placement Summary: {order_summary_str}")
        print(f"  Error Message: {step_state.get('error_message')}")
        print("-" * (len(step_name) + 24))
        # --- End State Printing ---

        final_state = step_state # Capture the last state

        # Halt if no items found after ingestion
        if step_name == 'data_ingestion' and isinstance(items_proc, list) and not items_proc:
             print("\n--- Halting workflow: No items found to process after ingestion. ---")
             break

    print("\n--- Graph Execution Finished ---")

    # --- Generate Human-Readable Summary Report ---
    print("\n\n========================================")
    print("   AI Agent Workflow Summary Report")
    print("========================================")

    if final_state:
        # --- Overall Status ---
        error_msg = final_state.get("error_message")
        ingestion_status = final_state.get("data_ingestion_status", "Unknown")
        items_processed_flag = isinstance(final_state.get('items_to_process'), list) and len(final_state.get('items_to_process')) > 0

        if error_msg and ingestion_status != "Success":
            print("** Overall Status: CRITICAL ERROR **")
        elif error_msg:
             print("** Overall Status: Completed with Warnings/Errors **")
        elif not items_processed_flag and ingestion_status == "Success":
             print("** Overall Status: Completed (No items required processing) **")
        else:
            print("** Overall Status: Completed Successfully **")

        # --- Data & Setup ---
        print("\n** Data & Setup **")
        print(f"- Data Ingestion Status: {ingestion_status}")
        items_count = len(final_state.get('items_to_process', [])) if isinstance(final_state.get('items_to_process'), list) else 0
        print(f"- Items Selected for Processing: {items_count}")

        # --- Analysis Results ---
        if items_count > 0: # Only show analysis if items were processed
            print("\n** Analysis Results **")
            fc_res = final_state.get('forecast_results', {})
            fc_proc = fc_res.get('processed_count', 0)
            fc_succ = fc_res.get('success_count', 0)
            print(f"- Demand Forecasting: Processed={fc_proc}, Succeeded={fc_succ}")

            inv_res = final_state.get('inventory_status', {})
            inv_proc = inv_res.get('summary', {}).get('processed_count', 0)
            # Count specific flags found during monitoring
            low_stock_count = 0
            excess_stock_count = 0
            near_expiry_count = 0
            expired_count = 0
            if isinstance(inv_res, dict):
                for key, value in inv_res.items():
                    if key != 'summary' and isinstance(value, dict):
                        flags = value.get('flags', [])
                        if 'LOW_STOCK' in flags: low_stock_count += 1
                        if 'EXCESS_STOCK' in flags: excess_stock_count += 1
                        if 'NEAR_EXPIRY' in flags: near_expiry_count += 1
                        if 'EXPIRED' in flags: expired_count += 1
            print(f"- Inventory Monitoring: Processed={inv_proc}")
            print(f"    - Items Flagged Low Stock:   {low_stock_count}")
            print(f"    - Items Flagged Excess Stock: {excess_stock_count}")
            print(f"    - Items Flagged Near Expiry: {near_expiry_count}")
            print(f"    - Items Flagged Expired:     {expired_count}")

            # --- Action Summary ---
            print("\n** Actions Proposed / Taken **")
            rep_res = final_state.get('replenishment_proposals', {})
            rep_prop_count = rep_res.get('proposal_count', 0)
            print(f"- Replenishment Orders Proposed: {rep_prop_count}")
            if rep_prop_count > 0:
                print("    Examples:")
                for order in rep_res.get('proposed_orders', [])[:3]: # Show first 3 examples
                    print(f"      - P:{order.get('ProductID')} S:{order.get('StoreID')} Qty:{order.get('QuantityOrdered')}")

            sup_res = final_state.get('order_placement_results', {})
            sup_placed_count = sup_res.get('placed_count', 0)
            print(f"- Orders Placed (Simulated): {sup_placed_count}")
            if sup_placed_count > 0:
                 print("    Examples:")
                 for order in sup_res.get('placed_orders', [])[:3]: # Show first 3 examples
                     print(f"      - ID:{order.get('OrderID')} P:{order.get('ProductID')} S:{order.get('StoreID')} Qty:{order.get('QuantityOrdered')} ETA:{order.get('ExpectedDeliveryDate')}")

            prc_res = final_state.get('pricing_proposals', {})
            prc_prop_count = prc_res.get('proposal_count', 0)
            print(f"- Pricing Actions Proposed: {prc_prop_count}")
            if prc_prop_count > 0:
                 print("    Examples:")
                 for action in prc_res.get('proposed_actions', [])[:3]: # Show first 3 examples
                      print(f"      - P:{action.get('ProductID')} S:{action.get('StoreID')} New Price:{action.get('RecommendedPrice')} (Reason: {action.get('Reason')})")

        # --- Errors / Warnings ---
        if error_msg:
            print("\n** Warnings / Errors Logged **")
            # Print only the message part for clarity
            error_lines = str(error_msg).split('\n')
            print(f"- {error_lines[0]}") # Print first line or the whole thing if one line
            if len(error_lines) > 1:
                 print("  (Check console log for full details if needed)")


    else:
        print("\n** Workflow execution did not produce a final state **")

    print("\n========================================")